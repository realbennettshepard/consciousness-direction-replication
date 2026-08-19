"""Render RESULTS.md to RESULTS.docx.

RESULTS.md IS THE SOURCE OF TRUTH. An earlier version of this script kept its own
hardcoded copy of the prose, which silently went stale the moment RESULTS.md was
edited -- the two documents disagreed and nothing complained. This one parses the
markdown, so they cannot diverge.

Supports the subset RESULTS.md actually uses: h1/h2/h3, paragraphs, pipe tables,
ordered and unordered lists, horizontal rules, and inline **bold** / *italic* /
`code` / ~~strike~~ / [links](url).

Usage:  python3 make_docx.py [input.md] [output.docx]
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SRC = Path(sys.argv[1] if len(sys.argv) > 1 else "RESULTS.md")
DST = Path(sys.argv[2] if len(sys.argv) > 2 else "RESULTS.docx")

FONT, MONO = "Calibri", "Consolas"
NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x59, 0x59, 0x59)
LINK = RGBColor(0x1F, 0x4E, 0x79)
HDR_FILL, ALT_FILL = "1F3864", "EEF2F8"
CONTENT_IN = 6.7

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.62)
sec.left_margin = sec.right_margin = Inches(0.9)
st = doc.styles["Normal"]
st.font.name, st.font.size = FONT, Pt(9.5)
st.paragraph_format.space_after = Pt(4)

# ------------------------------------------------------------------ inline spans
TOKEN = re.compile(
    r"\*\*(?P<b>[^*]+?)\*\*"                    # bold
    r"|(?<!\*)\*(?P<i>[^*]+?)\*(?!\*)"          # italic
    r"|`(?P<c>[^`]+?)`"                         # code
    r"|~~(?P<s>[^~]+?)~~"                       # strikethrough
    r"|\[(?P<lt>[^\]]+?)\]\((?P<lu>[^)]+?)\)")  # link


def add_runs(par, text, size=9.5):
    pos = 0
    for m in TOKEN.finditer(text):
        if m.start() > pos:
            r = par.add_run(text[pos:m.start()])
            r.font.name, r.font.size = FONT, Pt(size)
        if m.group("b") is not None:
            r = par.add_run(m.group("b")); r.bold = True; r.font.name = FONT
        elif m.group("i") is not None:
            r = par.add_run(m.group("i")); r.italic = True; r.font.name = FONT
        elif m.group("c") is not None:
            r = par.add_run(m.group("c")); r.font.name = MONO
        elif m.group("s") is not None:
            r = par.add_run(m.group("s")); r.font.strike = True
            r.font.name = FONT; r.font.color.rgb = GREY
        else:
            r = par.add_run(m.group("lt")); r.font.name = FONT
            r.font.color.rgb = LINK; r.underline = True
        r.font.size = Pt(size)
        pos = m.end()
    if pos < len(text):
        r = par.add_run(text[pos:])
        r.font.name, r.font.size = FONT, Pt(size)


def shade(cell, fill):
    e = OxmlElement("w:shd")
    e.set(qn("w:val"), "clear"); e.set(qn("w:color"), "auto"); e.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(e)


def heading(text, level):
    # Use Word's built-in Heading styles so the document is navigable (nav pane, TOC,
    # collapsible sections). Run-level font settings below override the styles' own
    # look, so the visual design is unchanged -- only the outline structure is added.
    try:
        p = doc.add_paragraph(style=f"Heading {level}")
    except KeyError:
        p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt({1: 2, 2: 11, 3: 8}[level])
    p.paragraph_format.space_after = Pt({1: 3, 2: 5, 3: 3}[level])
    r = p.add_run(re.sub(r"[*`]", "", text))
    r.font.name, r.bold = FONT, True
    r.font.size = Pt({1: 15, 2: 11, 3: 10}[level])
    r.font.color.rgb = NAVY if level > 1 else RGBColor(0, 0, 0)


def rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(7)
    bdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6")
    b.set(qn("w:color"), "BFBFBF"); b.set(qn("w:space"), "1")
    bdr.append(b); p._p.get_or_add_pPr().append(bdr)


def emit_table(rows):
    header, body = rows[0], rows[2:]          # rows[1] is the |---| separator
    ncol = len(header)
    prose = (max((len(c) for r in body for c in r[:1]), default=0) > 24)
    w0 = 2.05 if prose else CONTENT_IN / ncol
    widths = ([w0] + [(CONTENT_IN - w0) / (ncol - 1)] * (ncol - 1)) if ncol > 1 else [CONTENT_IN]
    t = doc.add_table(rows=1, cols=ncol)
    t.alignment, t.autofit = WD_TABLE_ALIGNMENT.LEFT, False
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    blank_hdr = all(not c.strip() for c in header)
    for i, cap in enumerate(header):
        c = t.rows[0].cells[i]
        if not blank_hdr:
            shade(c, HDR_FILL)
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
        if i and not prose:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(re.sub(r"[*`]", "", cap))
        r.font.name, r.font.size, r.bold = FONT, Pt(9), True
        if not blank_hdr:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(body):
        cells = t.add_row().cells
        for i, val in enumerate(row[:ncol]):
            if ri % 2:
                shade(cells[i], ALT_FILL)
            p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            if i and not prose:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(p, val, size=9)
    for row in t.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for i, c in enumerate(row.cells):
            c.width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def emit_list(items, ordered):
    """Explicit numerals so every list restarts at 1 -- python-docx's List Number
    style shares one numbering instance across lists and would continue counting."""
    for n, txt in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.30)
        p.paragraph_format.first_line_indent = Inches(-0.30)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{n}.  " if ordered else "•  ")
        r.font.name, r.font.size, r.bold = FONT, Pt(9.5), ordered
        add_runs(p, txt)


# ------------------------------------------------------------------ parse & emit
para, tbl, lst = [], [], []
lst_ord = False


def flush():
    global para, tbl, lst
    if para:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
        add_runs(p, " ".join(para))
        para = []
    if tbl:
        emit_table(tbl); tbl = []
    if lst:
        emit_list(lst, lst_ord); lst = []


for ln in SRC.read_text().split("\n"):
    s = ln.strip()
    if not s:
        flush()
    elif s.startswith("#"):
        flush()
        lvl = len(s) - len(s.lstrip("#"))
        heading(s.lstrip("#").strip(), min(lvl, 3))
    elif re.fullmatch(r"-{3,}", s):
        flush(); rule()
    elif s.startswith("|"):
        if para or lst:
            flush()
        tbl.append([c.strip() for c in s.strip("|").split("|")])
    elif re.match(r"^\d+\.\s", s):
        if para or tbl:
            flush()
        lst_ord = True
        lst.append(re.sub(r"^\d+\.\s+", "", s))
    elif s.startswith("- ") or s.startswith("* "):
        if para or tbl:
            flush()
        lst_ord = False
        lst.append(s[2:])
    elif lst and ln.startswith("   "):          # continuation of a list item
        lst[-1] += " " + s
    else:
        if tbl or lst:
            flush()
        para.append(s)
flush()

doc.save(DST)
print(f"rendered {SRC} -> {DST}")
